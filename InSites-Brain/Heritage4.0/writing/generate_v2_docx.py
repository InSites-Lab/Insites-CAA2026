"""
Generate Heritage4_0-Draft-v2.docx from the v2 markdown draft,
using Heritage4_0-Draft-v1.docx as the style template (Springer CCIS).

Key: all subsections use heading3 (no spacing) instead of heading2 (big spacing).
"""

import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from pathlib import Path

BASE = Path(r"c:\Users\user\Documents\InSItes-Workshops-26\CAA Workshop\InSites-Brain\Heritage4.0\writing")
TEMPLATE = BASE / "Heritage4_0-Draft-v1.docx"
MD_SOURCE = BASE / "Heritage4_0-Draft-v2.md"
OUTPUT = BASE / "Heritage4_0-Draft-v2.docx"


def read_md():
    """Parse the markdown into structured blocks."""
    with open(MD_SOURCE, "r", encoding="utf-8") as f:
        content = f.read()

    # Strip HTML comments
    content = re.sub(r"<!--.*?-->", "", content, flags=re.DOTALL)

    blocks = []
    lines = content.split("\n")
    i = 0
    in_references = False
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines and horizontal rules
        if not line or line == "---":
            i += 1
            continue

        # Title (# heading)
        if line.startswith("# ") and not line.startswith("## "):
            blocks.append(("title", line[2:].strip()))
            i += 1
            continue

        # Section heading (## N. Title)
        if line.startswith("## "):
            heading_text = line[3:].strip()
            if heading_text == "References":
                in_references = True
            blocks.append(("heading1", heading_text))
            i += 1
            continue

        # Subsection heading (### N.N Title)
        if line.startswith("### "):
            blocks.append(("heading3", line[4:].strip()))
            i += 1
            continue

        # Keywords line
        if line.startswith("**Keywords:**"):
            blocks.append(("keywords", line.replace("**Keywords:**", "Keywords:").strip()))
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1] and "---" in lines[i + 1]:
            # Parse markdown table
            header = [c.strip() for c in line.split("|")[1:-1]]
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].split("|")[1:-1]]
                rows.append(row)
                i += 1
            blocks.append(("table", (header, rows)))
            continue

        # Table caption (bold starting with "Table N.")
        if line.startswith("**Table "):
            caption = line.replace("**", "")
            blocks.append(("tablecaption", caption))
            i += 1
            continue

        # Check if we're in References section (each line = one reference)
        if in_references:
            blocks.append(("reference", line.strip()))
            i += 1
            continue

        # Regular paragraph (collect continuation lines)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("---"):
            if lines[i].startswith("## ") or lines[i].startswith("### ") or lines[i].startswith("# ") or lines[i].startswith("**Table ") or lines[i].startswith("**Keywords:"):
                break
            if lines[i] == "---":
                break
            # Check if next line starts a table
            if "|" in lines[i] and i + 1 < len(lines) and "|" in lines[i + 1] and "---" in lines[i + 1]:
                break
            para_lines.append(lines[i].rstrip())
            i += 1

        para_text = " ".join(para_lines)
        # Clean up markdown formatting for docx
        para_text = para_text.strip()
        if para_text:
            blocks.append(("paragraph", para_text))

    return blocks


def add_formatted_paragraph(doc, text, style_name):
    """Add a paragraph with inline bold/italic formatting from markdown."""
    p = doc.add_paragraph(style=style_name)

    # Split on bold and italic markers
    # Process **bold** and *italic* markers
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            # Handle citation brackets and special chars
            p.add_run(part)

    return p


def build_docx(blocks):
    """Build the docx from parsed blocks using the template styles."""
    # Open template to get styles
    doc = Document(str(TEMPLATE))

    # Remove all existing content (paragraphs + tables)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    for t in doc.tables:
        t._element.getparent().remove(t._element)

    first_para = True
    current_section = ""
    for block_type, content in blocks:
        if block_type == "title":
            doc.add_paragraph(content, style="papertitle")

        elif block_type == "heading1":
            doc.add_paragraph(content, style="heading1")
            current_section = content.strip()
            first_para = True  # next para gets p1a style

        elif block_type == "heading3":
            doc.add_paragraph(content, style="Heading 3")
            first_para = True

        elif block_type == "keywords":
            doc.add_paragraph(content, style="keywords")

        elif block_type == "tablecaption":
            doc.add_paragraph(content, style="tablecaption")

        elif block_type == "paragraph":
            # Abstract section gets abstract style
            if current_section == "Abstract":
                # Prefix with "Abstract. " in the style
                add_formatted_paragraph(doc, "Abstract. " + content, "abstract")
                current_section = ""  # only first para
            else:
                # First paragraph after heading uses p1a (no indent), rest use Normal
                style = "p1a" if first_para else "Normal"
                add_formatted_paragraph(doc, content, style)
                first_para = False

        elif block_type == "reference":
            doc.add_paragraph(content, style="referenceitem")

        elif block_type == "table":
            header, rows = content
            num_cols = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.style = "Normal Table"

            # Header row
            for j, cell_text in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(8)

            # Data rows
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j < num_cols:
                        cell = table.rows[i + 1].cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        run.font.size = Pt(8)

            first_para = True

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    blocks = read_md()
    print(f"Parsed {len(blocks)} blocks:")
    for btype, content in blocks:
        preview = str(content)[:60] if not isinstance(content, tuple) else f"table: {len(content[1])} rows"
        print(f"  {btype}: {preview}")
    print()
    build_docx(blocks)
